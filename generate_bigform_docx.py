import json
from docx import Document

with open('db/forms.json') as f:
    data = json.load(f)

sections = json.loads(data[1]['fields']['sections'])
doc = Document()
doc.add_heading('Big Form Placeholders', 0)

for section in sections:
    doc.add_heading(section['name'], level=1)
    for field in section['fields']:
        doc.add_paragraph('{{' + field['key'] + '}}')

doc.save('templates/bigform-2026-01-25.docx')
print('DOCX generated successfully.')
