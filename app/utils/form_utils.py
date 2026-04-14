import json
import os
import logging
import uuid
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt
from typing import Any, Dict, List, Tuple, Set
from app.exceptions import StorageError, TemplateFillingError, FileOperationError, InvalidFieldsError

# File to store forms
DB_DIR = "db"
FORMS_FILE = os.path.join(DB_DIR, "forms.json")
FORM_SUBMISSIONS_DB = os.path.join(DB_DIR, "form_submissions.json")
SAVED_FORM_SUBMISSIONS_DB = os.path.join(DB_DIR, "saved_form_submissions.json")
GENERATED_DIR = "generated"

# Configure logging
logger = logging.getLogger(__name__)

# Ensure directories exist
os.makedirs(DB_DIR, exist_ok=True)
os.makedirs(GENERATED_DIR, exist_ok=True)

def load_forms():
    """Load forms from JSON file with error handling."""
    try:
        if os.path.exists(FORMS_FILE):
            with open(FORMS_FILE, "r") as f:
                data = json.load(f)
                return data
        return []
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error while reading {FORMS_FILE}: {str(e)}")
        raise StorageError(f"Corrupted forms file: {str(e)}")
    except IOError as e:
        logger.error(f"IO error while reading {FORMS_FILE}: {str(e)}")
        raise StorageError(f"Cannot read forms file: {str(e)}")
    except Exception as e:
        logger.error(f"Unexpected error while loading forms: {str(e)}")
        raise StorageError(f"Unexpected error loading forms: {str(e)}")

def save_forms(forms):
    """Save forms to JSON file with error handling."""
    try:
        with open(FORMS_FILE, "w") as f:
            json.dump(forms, f, indent=4)
    except IOError as e:
        logger.error(f"IO error while writing to {FORMS_FILE}: {str(e)}")
        raise StorageError(f"Cannot save forms file: {str(e)}")
    except TypeError as e:
        logger.error(f"Type error while serializing forms: {str(e)}")
        raise StorageError(f"Forms contain non-serializable data: {str(e)}")
    except Exception as e:
        logger.error(f"Unexpected error while saving forms: {str(e)}")
        raise StorageError(f"Unexpected error saving forms: {str(e)}")

def load_form_submissions():
    """Load form submissions database with error handling."""
    try:
        if os.path.exists(FORM_SUBMISSIONS_DB):
            with open(FORM_SUBMISSIONS_DB, "r") as f:
                data = json.load(f)
                return data
        return {}
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error while reading {FORM_SUBMISSIONS_DB}: {str(e)}")
        raise StorageError(f"Corrupted form submissions database: {str(e)}")
    except IOError as e:
        logger.error(f"IO error while reading {FORM_SUBMISSIONS_DB}: {str(e)}")
        raise StorageError(f"Cannot read form submissions database: {str(e)}")
    except Exception as e:
        logger.error(f"Unexpected error while loading form submissions: {str(e)}")
        raise StorageError(f"Unexpected error loading form submissions: {str(e)}")

def save_form_submissions(form_submissions: Dict):
    """Save form submissions database with error handling."""
    try:
        with open(FORM_SUBMISSIONS_DB, "w") as f:
            json.dump(form_submissions, f, indent=4)
    except IOError as e:
        logger.error(f"IO error while writing to {FORM_SUBMISSIONS_DB}: {str(e)}")
        raise StorageError(f"Cannot save form submissions database: {str(e)}")
    except TypeError as e:
        logger.error(f"Type error while serializing form submissions: {str(e)}")
        raise StorageError(f"Form submissions contain non-serializable data: {str(e)}")
    except Exception as e:
        logger.error(f"Unexpected error while saving form submissions: {str(e)}")
        raise StorageError(f"Unexpected error saving form submissions: {str(e)}")


def load_saved_form_submissions() -> Dict:
    """Load saved form submissions database with error handling."""
    try:
        if os.path.exists(SAVED_FORM_SUBMISSIONS_DB):
            with open(SAVED_FORM_SUBMISSIONS_DB, "r") as f:
                return json.load(f)
        return {}
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error while reading {SAVED_FORM_SUBMISSIONS_DB}: {str(e)}")
        raise StorageError(f"Corrupted saved form submissions database: {str(e)}")
    except IOError as e:
        logger.error(f"IO error while reading {SAVED_FORM_SUBMISSIONS_DB}: {str(e)}")
        raise StorageError(f"Cannot read saved form submissions database: {str(e)}")
    except Exception as e:
        logger.error(f"Unexpected error while loading saved form submissions: {str(e)}")
        raise StorageError(f"Unexpected error loading saved form submissions: {str(e)}")


def save_saved_form_submissions(saved_submissions: Dict) -> None:
    """Save saved form submissions database with error handling."""
    try:
        with open(SAVED_FORM_SUBMISSIONS_DB, "w") as f:
            json.dump(saved_submissions, f, indent=4)
    except IOError as e:
        logger.error(f"IO error while writing to {SAVED_FORM_SUBMISSIONS_DB}: {str(e)}")
        raise StorageError(f"Cannot save saved form submissions database: {str(e)}")
    except TypeError as e:
        logger.error(f"Type error while serializing saved form submissions: {str(e)}")
        raise StorageError(f"Saved form submissions contain non-serializable data: {str(e)}")
    except Exception as e:
        logger.error(f"Unexpected error while saving saved form submissions: {str(e)}")
        raise StorageError(f"Unexpected error saving saved form submissions: {str(e)}")


def add_saved_form_submission(
    form_id: str,
    submission_id: str,
    values_used: Dict[str, Any],
    reference_text: str
) -> str:
    """Store values_used for a generated file under a form and return the submission ID."""
    try:
        saved_submissions = load_saved_form_submissions()

        if form_id not in saved_submissions:
            saved_submissions[form_id] = []

        saved_submission_id = str(uuid.uuid4())
        submission_record = {
            "saved_submission_id": saved_submission_id,
            "submission_id": submission_id,
            "reference_text": reference_text,
            "saved_at": datetime.now().isoformat(),
            "values_used": values_used,
        }

        saved_submissions[form_id].append(submission_record)
        save_saved_form_submissions(saved_submissions)
        return saved_submission_id
    except StorageError:
        raise
    except Exception as e:
        logger.error(f"Error adding saved form submission for form {form_id}, submission {submission_id}: {str(e)}")
        raise StorageError(f"Failed to add saved form submission: {str(e)}")


def delete_saved_form_submission(form_id: str, saved_submission_id: str) -> int:
    """Delete saved submission records by form_id and saved_submission_id and return delete count."""
    try:
        saved_submissions = load_saved_form_submissions()
        form_entries = saved_submissions.get(form_id, [])

        if not form_entries:
            return 0

        remaining_entries = [entry for entry in form_entries if entry.get("saved_submission_id") != saved_submission_id]
        deleted_count = len(form_entries) - len(remaining_entries)

        if deleted_count == 0:
            return 0

        if remaining_entries:
            saved_submissions[form_id] = remaining_entries
        else:
            saved_submissions.pop(form_id, None)

        save_saved_form_submissions(saved_submissions)
        return deleted_count
    except StorageError:
        raise
    except Exception as e:
        logger.error(f"Error deleting saved form submission for form {form_id}, saved_submission {saved_submission_id}: {str(e)}")
        raise StorageError(f"Failed to delete saved form submission: {str(e)}")

def add_form_submission_file(form_id: str, file_path: str, values: Dict[str, str]) -> str:
    """
    Register a form submission file in the database.
    
    Args:
        form_id: The form ID
        file_path: The path to the generated file
        values: The values used to fill the form
    
    Returns:
        The form submission file ID
    """
    try:
        form_submissions = load_form_submissions()
        
        # Create form entry if it doesn't exist
        if form_id not in form_submissions:
            form_submissions[form_id] = []
        
        # Create submission record
        submission_id = str(uuid.uuid4())
        submission_record = {
            "submission_id": submission_id,
            "file_path": file_path,
            "filename": os.path.basename(file_path),
            "created_at": datetime.now().isoformat(),
            "values_used": values
        }
        
        form_submissions[form_id].append(submission_record)
        save_form_submissions(form_submissions)
        return submission_id
    
    except StorageError:
        raise
    except Exception as e:
        logger.error(f"Error registering form submission file: {str(e)}")
        raise StorageError(f"Failed to register form submission file: {str(e)}")
def extract_placeholders_from_document(doc_path: str) -> Set[str]:
    """
    Extract all placeholder field names from a DOCX document.
    
    Looks for placeholders in format: {{field_name}}
    
    Args:
        doc_path: Path to the DOCX file
    
    Returns:
        Set of placeholder field names found in the document
    """
    try:
        if not os.path.exists(doc_path):
            logger.error(f"Document file not found: {doc_path}")
            raise FileOperationError(f"Document file not found: {doc_path}")
        
        doc = Document(doc_path)
        placeholders = set()
        
        # Regex pattern to find {{field_name}}
        pattern = r'\{\{([a-zA-Z0-9_]+)\}\}'
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            matches = re.findall(pattern, paragraph.text)
            if matches:
                placeholders.update(matches)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = re.findall(pattern, cell.text)
                    if matches:
                        placeholders.update(matches)
        
        # Extract from headers
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                matches = re.findall(pattern, paragraph.text)
                if matches:
                    placeholders.update(matches)
            
            # Extract from footers
            for paragraph in section.footer.paragraphs:
                matches = re.findall(pattern, paragraph.text)
                if matches:
                    placeholders.update(matches)
        
        return placeholders
    
    except FileOperationError:
        raise
    except Exception as e:
        logger.error(f"Error extracting placeholders from {doc_path}: {str(e)}")
        raise FileOperationError(f"Failed to extract placeholders: {str(e)}")

def validate_sections_against_document(doc_path: str, form_data: List) -> Dict[str, any]:
    """
    Validate that all keys in form sections fields are defined as placeholders in the document.
    
    This function extracts "key" values from nested fields in sections.
    Form structure (direct array):
    [
        {
            "name": "section_name",
            "fields": [
                {"name": "field_name", "datatype": "string", "key": "field_key"},
                ...
            ]
        }
    ]
    
    Any "key" found in sections that is NOT in the document is considered an "unknown key".
    
    Args:
        doc_path: Path to the DOCX template file
        form_data: Direct array of sections with nested "fields"
    
    Returns:
        Dictionary with validation results:
        {
            "valid": bool,
            "unknown_keys": List[str],  # Field keys not found in document
            "matched_keys": List[str],   # Field keys found in document
            "total_field_keys": int
        }
    """
    try:
        # Extract field keys from form_data
        field_keys = set()
                
        if isinstance(form_data, list):
            for idx, section in enumerate(form_data):
                if isinstance(section, dict) and "fields" in section:
                    fields_array = section["fields"]
                    
                    if isinstance(fields_array, list):
                        for field_idx, field in enumerate(fields_array):
                            if isinstance(field, dict) and "key" in field:
                                key = field["key"]
                                field_keys.add(key)
                            else:
                                logger.warning(f"Section {idx}, Field {field_idx}: No 'key' field found")
                else:
                    logger.warning(f"Section {idx}: No 'fields' array found or not a dict")
        else:
            logger.warning("form_data is not a list")
        
        # Extract placeholders from document
        placeholders = extract_placeholders_from_document(doc_path)
        
        # Find unknown keys (field keys not in document)
        unknown_keys = field_keys - placeholders
        matched_keys = field_keys & placeholders
        
        is_valid = len(unknown_keys) == 0
        
        result = {
            "valid": is_valid,
            "unknown_keys": sorted(list(unknown_keys)),
            "matched_keys": sorted(list(matched_keys)),
            "total_field_keys": len(field_keys)
        }
        
        return result
    
    except FileOperationError:
        raise
    except Exception as e:
        logger.error(f"Error validating sections against document: {str(e)}", exc_info=True)
        raise FileOperationError(f"Failed to validate sections: {str(e)}")

def validate_form_fields(doc_path: str, form_fields: Dict[str, str]) -> Dict[str, any]:
    """
    Validate that all placeholders in the document match the form fields.
    
    Args:
        doc_path: Path to the DOCX template file
        form_fields: Dictionary of form fields (keys are field names)
    
    Returns:
        Dictionary with validation results:
        {
            "valid": bool,
            "missing_in_fields": List[str],  # Placeholders not in form fields
            "unused_fields": List[str],       # Form fields not in placeholders
            "matched_fields": List[str],      # Fields that match
            "total_placeholders": int,
            "total_fields": int
        }
    """
    try:
        # Extract placeholders from document
        placeholders = extract_placeholders_from_document(doc_path)
        
        # Get form field keys
        field_keys = set(form_fields.keys())
        
        # Find differences
        missing_in_fields = placeholders - field_keys
        unused_fields = field_keys - placeholders
        matched_fields = placeholders & field_keys
        
        is_valid = len(missing_in_fields) == 0 and len(unused_fields) == 0
        
        result = {
            "valid": is_valid,
            "missing_in_fields": sorted(list(missing_in_fields)),
            "unused_fields": sorted(list(unused_fields)),
            "matched_fields": sorted(list(matched_fields)),
            "total_placeholders": len(placeholders),
            "total_fields": len(field_keys)
        }
        
        return result
    
    except FileOperationError:
        raise
    except Exception as e:
        logger.error(f"Error validating form fields: {str(e)}")
        raise FileOperationError(f"Failed to validate form fields: {str(e)}")
def get_form_submissions(form_id: str) -> List[Dict]:
    """Get all form submissions for a form."""
    try:
        form_submissions = load_form_submissions()
        files = form_submissions.get(form_id, [])
        return files
    except StorageError:
        raise
    except Exception as e:
        logger.error(f"Error retrieving form submissions for form {form_id}: {str(e)}")
        raise StorageError(f"Failed to retrieve form submissions: {str(e)}")

def _apply_run_style(run, style: Dict[str, Any]) -> None:
    """Apply supported style attributes to a DOCX run."""

    # Set font family (accepts either font_name or font_family key).
    font_name = style.get("font_name") or style.get("font_family")
    if font_name:
        run.font.name = str(font_name)

    # Set font size in points.
    font_size = style.get("font_size", style.get("size"))
    if font_size is not None:
        try:
            run.font.size = Pt(float(font_size))
        except (TypeError, ValueError):
            # Keep document generation running even if size input is invalid.
            logger.warning(f"Invalid font size provided: {font_size}")

def _replace_placeholders_in_paragraph(paragraph, values: Dict[str, Any], default_style: Dict[str, Any] = None) -> None:
    """Replace placeholders in one paragraph and apply optional run styles."""
    # Use empty style when no defaults are passed.
    default_style = default_style or {}
    # Read full paragraph text (combined from all runs).
    original_text = paragraph.text
    if not original_text:
        return

    # Match placeholders like {{field_name}}.
    pattern = r"\{\{([a-zA-Z0-9_]+)\}\}"
    matches = list(re.finditer(pattern, original_text))
    # Nothing to replace in this paragraph.
    if not matches:
        return

    # segments stores final paragraph parts as (text, style).
    segments: List[Tuple[str, Dict[str, Any]]] = []
    last_pos = 0
    has_replacement = False

    for match in matches:
        # Extract key inside braces, e.g., {{name}} -> name.
        key = match.group(1)
        # Skip placeholders that are not present in values.
        if key not in values:
            continue

        # Keep normal text before the placeholder unchanged.
        if match.start() > last_pos:
            segments.append((original_text[last_pos:match.start()], {}))

        # Replace placeholder with provided value and apply default style.
        replacement_text = str(values[key])
        segments.append((replacement_text, dict(default_style)))
        last_pos = match.end()
        has_replacement = True

    # If no known keys were found, keep paragraph as-is.
    if not has_replacement:
        return

    # Keep remaining text after the last placeholder.
    if last_pos < len(original_text):
        segments.append((original_text[last_pos:], {}))

    # Remove old runs so we can rebuild paragraph content cleanly.
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)

    # Add new runs back in order and apply style only to replacement parts.
    for text, style in segments:
        if text == "":
            continue
        new_run = paragraph.add_run(text)
        if style:
            _apply_run_style(new_run, style)

def fill_template(
    template_path: str,
    values: Dict[str, Any],
    form_id: str = None,
    form_name: str = None,
    font_family: str = None,
    font_size: float = None,
) -> Tuple[str, str]:
    """
    Fill template with values, replacing placeholders in document.
    
    Args:
        template_path: Path to the template DOCX file
        values: Dictionary of values to fill
        form_id: Form ID for tracking generated files
        form_name: Form name for the generated filename
        font_family: Optional font family applied to all inserted values
        font_size: Optional font size (in points) applied to all inserted values
    
    Returns:
        Tuple of (filled_file_path, submission_id) where submission_id is used for tracking
    """
    try:
        if not os.path.exists(template_path):
            logger.error(f"Template file not found: {template_path}")
            raise FileOperationError(f"Template file not found: {template_path}")
        
        doc = Document(template_path)
        default_style = {}
        if font_family:
            default_style["font_family"] = font_family
        if font_size is not None:
            default_style["font_size"] = font_size
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            _replace_placeholders_in_paragraph(paragraph, values, default_style)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_placeholders_in_paragraph(paragraph, values, default_style)
        
        # Replace in headers and footers
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                _replace_placeholders_in_paragraph(paragraph, values, default_style)
            for paragraph in section.footer.paragraphs:
                _replace_placeholders_in_paragraph(paragraph, values, default_style)
        
        # Generate filename with form name and timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Use form_name if provided, otherwise use form_id
        file_prefix = form_name or form_id or "form"
        # Sanitize filename (remove special characters)
        file_prefix = "".join(c for c in file_prefix if c.isalnum() or c in ('-', '_')).rstrip()
        filled_filename = f"{file_prefix}_{timestamp}.docx"
        filled_path = os.path.join(GENERATED_DIR, filled_filename)
        
        # Create generated directory if it doesn't exist
        os.makedirs(GENERATED_DIR, exist_ok=True)
        
        doc.save(filled_path)
        
        # Register the generated file if form_id is provided
        submission_id = None
        if form_id:
            try:
                submission_id = add_form_submission_file(form_id, filled_path, values)
            except StorageError as e:
                logger.warning(f"Could not register generated file, but template was filled: {str(e)}")
                submission_id = str(uuid.uuid4())
        else:
            submission_id = str(uuid.uuid4())
        
        return filled_path, submission_id
        
    except FileOperationError:
        raise
    except Exception as e:
        logger.error(f"Error filling template {template_path}: {str(e)}")
        raise TemplateFillingError(f"Failed to fill template: {str(e)}")
